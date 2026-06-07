import re
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# 中文字号 → 半磅值 (w:sz 用的是半磅)
_CN_SIZES = {
    "初号": 84, "小初": 72, "一号": 52, "小一": 48,
    "二号": 44, "小二": 36, "三号": 32, "小三": 30,
    "四号": 28, "小四": 24, "五号": 21, "小五": 18,
}


class DocFormatter:
    def __init__(self, docx_path):
        self.doc = Document(docx_path)

    # ── extract structure for LLM ─────────────────────────────────────

    def extract_paragraphs(self, max_chars=80):
        """提取每个段落的索引、样式、文字，供 LLM 阅读理解"""
        result = []
        for i, p in enumerate(self.doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue
            result.append({
                "idx": i,
                "style": p.style.name,
                "text": text[:max_chars] + ("..." if len(text) > max_chars else ""),
            })
        return result

    # ── apply per-paragraph formatting from LLM ───────────────────────

    def execute(self, para_formats):
        """
        para_formats 格式:
        [
          {"idx": 0, "font": "黑体", "size_pt": 16, "bold": true, "align": "center"},
          {"idx": 2, "font": "宋体", "size_pt": 12, "indent_chars": 2, "line_spacing": 1.5},
          ...
        ]
        """
        for item in para_formats:
            idx = item.get("idx")
            if idx is None or idx >= len(self.doc.paragraphs):
                continue
            para = self.doc.paragraphs[idx]
            self._apply(para, item)

    def save_as(self, output_path):
        self.doc.save(output_path)

    # ── apply formatting to a paragraph ───────────────────────────────

    def _apply(self, para, fmt):
        # 字体 — 需要同时设置 ascii/hAnsi 和 eastAsia
        font_name = fmt.get("font")
        size_pt = fmt.get("size_pt")
        size_hint = fmt.get("size_hint")  # 中文字号名称如"小四"

        if size_hint and size_hint in _CN_SIZES:
            size_pt = _CN_SIZES[size_hint] / 2  # 半磅 → pt

        for run in para.runs:
            rpr = run._element.find(f"{{{NS}}}rPr")
            if rpr is None:
                rpr = etree.SubElement(run._element, f"{{{NS}}}rPr")

            # 设置字体
            if font_name or size_pt:
                rfonts = rpr.find(f"{{{NS}}}rFonts")
                if rfonts is None:
                    rfonts = etree.SubElement(rpr, f"{{{NS}}}rFonts")
                if font_name:
                    rfonts.set(f"{{{NS}}}ascii", font_name)
                    rfonts.set(f"{{{NS}}}hAnsi", font_name)
                    rfonts.set(f"{{{NS}}}eastAsia", font_name)
                    rfonts.set(f"{{{NS}}}cs", font_name)
                if size_pt:
                    sz_val = str(int(size_pt * 2))  # pt → 半磅
                    sz = rpr.find(f"{{{NS}}}sz")
                    if sz is None:
                        sz = etree.SubElement(rpr, f"{{{NS}}}sz")
                    sz.set(f"{{{NS}}}val", sz_val)
                    szCs = rpr.find(f"{{{NS}}}szCs")
                    if szCs is None:
                        szCs = etree.SubElement(rpr, f"{{{NS}}}szCs")
                    szCs.set(f"{{{NS}}}val", sz_val)

            # 加粗
            if "bold" in fmt:
                b = rpr.find(f"{{{NS}}}b")
                if b is None:
                    b = etree.SubElement(rpr, f"{{{NS}}}b")
                if fmt["bold"]:
                    b.attrib.pop(f"{{{NS}}}val", None)
                else:
                    b.set(f"{{{NS}}}val", "0")

            # 斜体
            if "italic" in fmt:
                i_elem = rpr.find(f"{{{NS}}}i")
                if i_elem is None:
                    i_elem = etree.SubElement(rpr, f"{{{NS}}}i")
                if fmt["italic"]:
                    i_elem.attrib.pop(f"{{{NS}}}val", None)
                else:
                    i_elem.set(f"{{{NS}}}val", "0")

            # 颜色
            if "color" in fmt:
                hex_str = fmt["color"].lstrip("#")
                color_elem = rpr.find(f"{{{NS}}}color")
                if color_elem is None:
                    color_elem = etree.SubElement(rpr, f"{{{NS}}}color")
                color_elem.set(f"{{{NS}}}val", hex_str)

        # 段落格式
        if "align" in fmt and fmt["align"] in ALIGN_MAP:
            para.alignment = ALIGN_MAP[fmt["align"]]

        pf = para.paragraph_format
        if "indent_chars" in fmt:
            # 用 w:firstLineChars 实现真正的"字符"缩进
            pPr = para._element.find(f"{{{NS}}}pPr")
            if pPr is None:
                pPr = etree.SubElement(para._element, f"{{{NS}}}pPr")
            ind = pPr.find(f"{{{NS}}}ind")
            if ind is None:
                ind = etree.SubElement(pPr, f"{{{NS}}}ind")
            ind.set(f"{{{NS}}}firstLineChars", str(int(float(fmt["indent_chars"]) * 100)))
            # 清除可能存在的 fixed 值
            ind.attrib.pop(f"{{{NS}}}firstLine", None)
        if "line_spacing" in fmt:
            pf.line_spacing = float(fmt["line_spacing"])
        if "line_spacing_pt" in fmt:
            pf.line_spacing = Pt(float(fmt["line_spacing_pt"]))
        if "space_before_pt" in fmt:
            pf.space_before = Pt(float(fmt["space_before_pt"]))
        if "space_after_pt" in fmt:
            pf.space_after = Pt(float(fmt["space_after_pt"]))
